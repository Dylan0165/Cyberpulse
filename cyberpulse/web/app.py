"""CyberPulse — FastAPI JSON API for the Tauri desktop app."""

import sys
from pathlib import Path
# Ensure /app (cyberpulse root) is in sys.path so 'reports', 'engine', etc. are importable
_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

import json
import logging
import secrets
import shutil
import threading
import time
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path

from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from starlette.middleware.base import BaseHTTPMiddleware

from config import Config
from engine.scanner import Scanner
from ai.analyzer import analyze_scan, analyze_scan_streaming
from reports.generator import generate_pdf

logger = logging.getLogger("cyberpulse.web")

# In-memory scan state: {scan_id: {"status": ..., "events": [...]}}
_active_scans: dict[str, dict] = {}
_scan_lock = threading.Lock()


@asynccontextmanager
async def lifespan(app: FastAPI):
    Config.ensure_dirs()
    for issue in Config.validate():
        logger.warning(issue)
    # Load persisted schedules and register APScheduler jobs
    loaded = _load_schedules_file()
    with _schedule_lock:
        _schedules.extend(loaded)
    for sched in loaded:
        _register_schedule_job(sched)
    logger.info("CyberPulse API gestart op %s:%s", Config.FLASK_HOST, Config.FLASK_PORT)
    try:
        yield
    finally:
        if _scan_scheduler:
            try:
                _scan_scheduler.shutdown(wait=False)
            except Exception:
                pass
        logger.info("CyberPulse API gestopt")


app = FastAPI(title="CyberPulse", version="3.0", lifespan=lifespan)


# CORS — allow Svelte dev, Tauri, and the FastAPI backend container
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "tauri://localhost",
        "https://tauri.localhost",
        "http://localhost:8000",
        "http://backend:8000",
        "http://localhost:3000",
        "http://frontend:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ──────────────────── API KEY AUTHENTICATION ────────────────────

# Paths that are always accessible without authentication
_AUTH_EXEMPT = {"/health", "/", "/docs", "/openapi.json", "/redoc"}


class ApiKeyMiddleware(BaseHTTPMiddleware):
    """Require X-API-Key header when Config.API_KEY is set.

    Returns 401 for any non-exempt path that lacks a valid key.
    When Config.API_KEY is empty the middleware is a no-op.
    """

    async def dispatch(self, request: Request, call_next):
        if not Config.API_KEY:
            return await call_next(request)
        if request.url.path in _AUTH_EXEMPT:
            return await call_next(request)
        key = request.headers.get("X-API-Key", "")
        # Use constant-time comparison to prevent timing attacks
        if not key or not secrets.compare_digest(key, Config.API_KEY):
            return JSONResponse(
                {"detail": "Ongeldige of ontbrekende X-API-Key header"},
                status_code=401,
            )
        return await call_next(request)


app.add_middleware(ApiKeyMiddleware)


# ─────────────────────── REQUEST MODELS ──────────────────────────

class ScanRequest(BaseModel):
    target: str
    scan_type: str = "quick"
    scan_mode: str = "blackbox"
    target_type: str = "web"
    modules: list[str] | None = None
    credentials: dict | None = None


class SettingsUpdate(BaseModel):
    deepseek_api_key: str | None = None
    deepseek_model: str | None = None
    deepseek_temperature: float | None = None
    nmap_timing: str | None = None
    max_threads: int | None = None
    scan_timeout: int | None = None
    wordlist_dir: str | None = None
    shodan_api_key: str | None = None
    abuseipdb_api_key: str | None = None
    virustotal_api_key: str | None = None
    hibp_api_key: str | None = None


# ──────────────────────────── API ROUTES ─────────────────────────



@app.get("/health")
async def health():
    return {"status": "ok", "version": "3.0"}


@app.get("/api/scans")
async def list_scans():
    """List recent scans."""
    return _get_recent_scans(limit=50)


@app.post("/api/scan/start")
async def start_scan(body: ScanRequest):
    """Start a new scan via JSON body."""
    target = body.target.strip()
    if not target:
        raise HTTPException(status_code=400, detail="Target is verplicht")
    # Normalize: strip scheme so modules can prepend http/https themselves
    if "://" in target:
        from urllib.parse import urlparse as _urlparse
        _p = _urlparse(target)
        target = _p.netloc + (_p.path.rstrip("/") if _p.path and _p.path != "/" else "")

    scan_type = body.scan_type if body.scan_type in ("quick", "standard", "full", "custom") else "quick"
    scan_mode = body.scan_mode if body.scan_mode in ("blackbox", "graybox", "whitebox") else "blackbox"
    target_type = body.target_type if body.target_type in (
        "web", "api", "spa", "mobile_ios", "mobile_android", "desktop", "network"
    ) else "web"

    # Determine modules
    if scan_type == "full":
        modules = Config.ALL_MODULES[:]
    elif scan_type == "custom" and body.modules:
        modules = body.modules
    elif scan_type == "standard":
        modules = Config.ALL_MODULES[:]
    else:
        modules = Config.QUICK_MODULES[:]

    # Add mode-specific modules
    if scan_mode == "graybox":
        for m in Config.GRAYBOX_MODULES:
            if m not in modules:
                modules.append(m)
    elif scan_mode == "whitebox":
        for m in Config.GRAYBOX_MODULES + Config.WHITEBOX_MODULES:
            if m not in modules:
                modules.append(m)

    # Add target-type specific modules
    if target_type == "mobile_android":
        for m in Config.MOBILE_ANDROID_MODULES:
            if m not in modules:
                modules.append(m)
    elif target_type == "mobile_ios":
        for m in Config.MOBILE_IOS_MODULES:
            if m not in modules:
                modules.append(m)
    elif target_type == "desktop":
        for m in Config.DESKTOP_MODULES:
            if m not in modules:
                modules.append(m)

    # Sanitize credentials
    credentials = {}
    if body.credentials:
        allowed_keys = {
            "web_username", "web_password", "web_login_url",
            "api_token", "api_header",
            "ssh_host", "ssh_port", "ssh_username", "ssh_password", "ssh_key",
            "db_host", "db_port", "db_username", "db_password", "db_name",
            "source_path", "network_range",
        }
        credentials = {k: v for k, v in body.credentials.items() if k in allowed_keys and v}

    scanner = Scanner(
        target=target, modules=modules, scan_type=scan_type,
        scan_mode=scan_mode, credentials=credentials, target_type=target_type,
    )
    scan_id = scanner.scan_id

    with _scan_lock:
        _active_scans[scan_id] = {"status": "running", "events": []}

    thread = threading.Thread(target=_run_scan_thread, args=(scanner, scan_id), daemon=True)
    thread.start()

    return {"scan_id": scan_id, "status": "started"}


@app.post("/api/scan/{scan_id}/stop")
async def stop_scan(scan_id: str):
    """Stop an active scan."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)

    with _scan_lock:
        scan_state = _active_scans.get(scan_id)
        if scan_state:
            scan_state["status"] = "stopped"
            scan_state["events"].append({"type": "scan_stopped", "message": "Scan handmatig gestopt"})

    return {"status": "stopped"}


@app.get("/api/scan/{scan_id}/stream")
async def scan_stream(scan_id: str):
    """SSE endpoint — streams scan events to the frontend."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)

    def event_generator():
        last_index = 0
        while True:
            with _scan_lock:
                scan_state = _active_scans.get(scan_id)

            if scan_state is None:
                scan_dir = Config.SCANS_DIR / scan_id
                if (scan_dir / "scan_data.json").exists():
                    yield f"data: {json.dumps({'type': 'redirect', 'url': f'/report/{scan_id}'})}\n\n"
                    return
                yield f"data: {json.dumps({'type': 'error', 'message': 'Scan niet gevonden'})}\n\n"
                return

            events = scan_state["events"]
            while last_index < len(events):
                yield f"data: {json.dumps(events[last_index])}\n\n"
                last_index += 1

            if scan_state["status"] in ("done", "stopped"):
                return

            time.sleep(0.3)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/api/scans/{scan_id}/report")
async def get_report(scan_id: str):
    """Get scan report data as JSON."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)

    scan_dir = Config.SCANS_DIR / scan_id
    scan_data_file = scan_dir / "scan_data.json"
    if not scan_data_file.exists():
        raise HTTPException(status_code=404, detail="Scan niet gevonden")

    with open(scan_data_file, "r", encoding="utf-8") as f:
        scan_data = json.load(f)

    analysis_file = scan_dir / "analysis.json"
    analysis = {}
    if analysis_file.exists():
        with open(analysis_file, "r", encoding="utf-8") as f:
            analysis = json.load(f)

    # Exploit correlation data
    correlation_file = scan_dir / "exploit_correlation.json"
    correlation = {}
    if correlation_file.exists():
        with open(correlation_file, "r", encoding="utf-8") as f:
            correlation = json.load(f)

    return {
        "scan_data": scan_data,
        "analysis": analysis,
        "correlation": correlation,
    }


@app.get("/api/scans/{scan_id}/pdf")
async def download_pdf_report(scan_id: str):
    """Download the PDF report."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)

    scan_dir = Config.SCANS_DIR / scan_id
    pdf_path = scan_dir / "report.pdf"

    if not pdf_path.exists():
        pdf_result = generate_pdf(scan_id)
        if pdf_result is None:
            raise HTTPException(status_code=404, detail="Rapport generatie mislukt")
        if pdf_result.suffix == ".html":
            return FileResponse(str(pdf_result), media_type="text/html",
                                filename=f"cyberpulse-{scan_id}.html")
        pdf_path = pdf_result

    return FileResponse(str(pdf_path), media_type="application/pdf",
                        filename=f"cyberpulse-{scan_id}.pdf")


@app.get("/api/scan/{scan_id}/analysis/stream")
async def analysis_stream(scan_id: str):
    """SSE endpoint — streams AI analysis to the frontend."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)

    scan_dir = Config.SCANS_DIR / scan_id
    scan_data_file = scan_dir / "scan_data.json"
    if not scan_data_file.exists():
        raise HTTPException(status_code=404)

    with open(scan_data_file, "r", encoding="utf-8") as f:
        scan_data = json.load(f)

    def analysis_generator():
        try:
            for chunk in analyze_scan_streaming(scan_data):
                yield f"data: {json.dumps({'type': 'chunk', 'text': chunk})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            logger.error("Analysis streaming failed: %s", e)
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(
        analysis_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ──────────────────────── SETTINGS ────────────────────────────

@app.get("/api/settings")
async def get_settings():
    """Get current settings (redacted)."""
    return {
        "deepseek_api_key": _redact(Config.DEEPSEEK_API_KEY),
        "deepseek_model": Config.DEEPSEEK_MODEL,
        "deepseek_temperature": Config.DEEPSEEK_TEMPERATURE,
        "nmap_timing": Config.NMAP_TIMING,
        "max_threads": Config.MAX_THREADS,
        "scan_timeout": Config.SCAN_TIMEOUT,
        "wordlist_dir": Config.WORDLIST_DIR,
        "shodan_api_key": _redact(Config.SHODAN_API_KEY),
        "abuseipdb_api_key": _redact(Config.ABUSEIPDB_API_KEY),
        "virustotal_api_key": _redact(Config.VIRUSTOTAL_API_KEY),
        "hibp_api_key": _redact(Config.HIBP_API_KEY),
    }


@app.put("/api/settings")
async def update_settings(body: SettingsUpdate):
    """Update settings in memory (non-persistent for security)."""
    if body.deepseek_api_key and not body.deepseek_api_key.startswith("****"):
        Config.DEEPSEEK_API_KEY = body.deepseek_api_key
    if body.deepseek_model:
        Config.DEEPSEEK_MODEL = body.deepseek_model
    if body.deepseek_temperature is not None:
        Config.DEEPSEEK_TEMPERATURE = body.deepseek_temperature
    if body.nmap_timing:
        Config.NMAP_TIMING = body.nmap_timing
    if body.max_threads is not None:
        Config.MAX_THREADS = body.max_threads
    if body.scan_timeout is not None:
        Config.SCAN_TIMEOUT = body.scan_timeout
    if body.wordlist_dir:
        Config.WORDLIST_DIR = body.wordlist_dir
    if body.shodan_api_key and not body.shodan_api_key.startswith("****"):
        Config.SHODAN_API_KEY = body.shodan_api_key
    if body.abuseipdb_api_key and not body.abuseipdb_api_key.startswith("****"):
        Config.ABUSEIPDB_API_KEY = body.abuseipdb_api_key
    if body.virustotal_api_key and not body.virustotal_api_key.startswith("****"):
        Config.VIRUSTOTAL_API_KEY = body.virustotal_api_key
    if body.hibp_api_key and not body.hibp_api_key.startswith("****"):
        Config.HIBP_API_KEY = body.hibp_api_key

    return {"status": "updated"}


@app.get("/api/tools/check")
async def check_tools():
    """Check which external tools are installed."""
    tools = [
        "nmap", "nikto", "gobuster", "ffuf", "sqlmap", "testssl.sh",
        "nuclei", "feroxbuster", "whatweb", "gitleaks", "msfconsole",
        "masscan", "wpscan",
    ]
    results = {}
    for tool in tools:
        results[tool] = shutil.which(tool) is not None
    return results


# ───────────────────── KALI TOOL ENDPOINTS ───────────────────────

class ToolScanRequest(BaseModel):
    target: str
    tools: list[str] | None = None
    profile: str | None = None
    scan_mode: str = "blackbox"   # blackbox | graybox | whitebox


# Extra tools unlocked per scan mode on top of the selected profile/tool list
_MODE_EXTRA_TOOLS: dict[str, list[str]] = {
    "graybox": ["crackmapexec", "evil-winrm", "impacket", "nikto", "wpscan"],
    "whitebox": ["lynis", "trivy", "grype", "chkrootkit", "binwalk", "exiftool"],
}


@app.get("/api/tools/available")
async def list_available_tools():
    """Return all registered Kali tools and their availability."""
    from tools.tool_runner import ToolRunner
    runner = ToolRunner()
    return runner.get_available_tools()


@app.get("/api/tools/profiles")
async def list_tool_profiles():
    """Return available tool scan profiles."""
    from tools.tool_runner import TOOL_SCAN_PROFILES
    return TOOL_SCAN_PROFILES


@app.post("/api/tools/scan")
async def start_tool_scan(body: ToolScanRequest):
    """Start a standalone Kali tool scan.

    scan_mode controls which extra tools are appended:
      - blackbox : only the selected tools / profile (external attacker view)
      - graybox  : adds internal/authenticated tools (crackmapexec, evil-winrm …)
      - whitebox : adds source/config analysis tools (lynis, trivy, grype …)
    """
    target = body.target.strip()
    if not target:
        raise HTTPException(status_code=400, detail="Target is verplicht")

    scan_mode = body.scan_mode if body.scan_mode in ("blackbox", "graybox", "whitebox") else "blackbox"

    from tools.tool_runner import ToolRunner, TOOL_SCAN_PROFILES
    import uuid

    # Determine base tool list
    if body.tools:
        tool_names = list(body.tools)
    elif body.profile and body.profile in TOOL_SCAN_PROFILES:
        tool_names = list(TOOL_SCAN_PROFILES[body.profile])
    else:
        tool_names = list(TOOL_SCAN_PROFILES.get("web_quick", ["nikto", "gobuster", "nuclei"]))

    # Append mode-specific extra tools (no duplicates)
    for extra in _MODE_EXTRA_TOOLS.get(scan_mode, []):
        if extra not in tool_names:
            tool_names.append(extra)

    scan_id = f"tool_{int(time.time())}_{uuid.uuid4().hex[:6]}"
    scan_dir = Config.SCANS_DIR / scan_id
    scan_dir.mkdir(parents=True, exist_ok=True)

    with _scan_lock:
        _active_scans[scan_id] = {"status": "running", "events": []}

    thread = threading.Thread(
        target=_run_tool_scan_thread,
        args=(scan_id, target, tool_names, scan_dir),
        daemon=True,
    )
    thread.start()

    return {"scan_id": scan_id, "status": "started", "tools": tool_names, "scan_mode": scan_mode}


def _run_tool_scan_thread(scan_id: str, target: str, tool_names: list[str], scan_dir: Path):
    """Run Kali tools in a background thread and collect events."""
    from tools.tool_runner import ToolRunner

    runner = ToolRunner()
    all_results = []
    try:
        gen = runner.run_tools(tool_names, target, scan_dir)
        for event in gen:
            with _scan_lock:
                scan_state = _active_scans.get(scan_id)
                if scan_state:
                    if scan_state["status"] == "stopped":
                        return
                    scan_state["events"].append(event)

        # Collect results (generator return value)
        # Since we can't get the return value from a generator that way,
        # we run tools individually for the final data
        for name in tool_names:
            result = runner.run_tool(name, target, scan_dir)
            all_results.append(result.to_dict())

        # Save tool scan data
        scan_data = {
            "scan_id": scan_id,
            "target": target,
            "scan_type": "toolscan",
            "started_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "tools_run": tool_names,
            "tool_results": all_results,
            "total_findings": sum(len(r.get("findings", [])) for r in all_results),
        }

        with open(scan_dir / "scan_data.json", "w", encoding="utf-8") as f:
            json.dump(scan_data, f, ensure_ascii=False, indent=2)

        _emit_event(scan_id, {"type": "all_done", "scan_id": scan_id})

    except Exception as e:
        logger.error("Tool scan thread failed: %s", e)
        _emit_event(scan_id, {"type": "error", "message": str(e)})
    finally:
        with _scan_lock:
            if scan_id in _active_scans:
                _active_scans[scan_id]["status"] = "done"


# ───────────────────────── BACKGROUND SCAN ───────────────────────

def _run_scan_thread(scanner: Scanner, scan_id: str):
    """Run the scan in a background thread and collect events."""
    try:
        for event in scanner.run():
            with _scan_lock:
                scan_state = _active_scans.get(scan_id)
                if scan_state:
                    if scan_state["status"] == "stopped":
                        _emit_event(scan_id, {"type": "scan_stopped"})
                        return
                    scan_state["events"].append(event)

        scan_dir = Config.SCANS_DIR / scan_id
        scan_data_file = scan_dir / "scan_data.json"
        if scan_data_file.exists():
            with open(scan_data_file, "r", encoding="utf-8") as f:
                scan_data = json.load(f)

            # Deduplicate findings before AI analysis
            try:
                from engine.findings_deduplicator import deduplicate_scan_data
                scan_data = deduplicate_scan_data(scan_data)
                with open(scan_data_file, "w", encoding="utf-8") as f:
                    json.dump(scan_data, f, ensure_ascii=False, indent=2)
            except Exception as e:
                logger.warning("Deduplicator failed (non-fatal): %s", e)

            _emit_event(scan_id, {"type": "analysis_start"})

            try:
                analysis = analyze_scan(scan_data)
                analysis_file = scan_dir / "analysis.json"
                with open(analysis_file, "w", encoding="utf-8") as f:
                    json.dump(analysis, f, ensure_ascii=False, indent=2)
                _emit_event(scan_id, {"type": "analysis_done"})
            except Exception as e:
                logger.error("AI analysis failed: %s", e)
                _emit_event(scan_id, {"type": "analysis_error", "message": str(e)})

            try:
                generate_pdf(scan_id)
                _emit_event(scan_id, {"type": "pdf_ready"})
            except Exception as e:
                logger.error("PDF generation failed: %s", e)

        _emit_event(scan_id, {"type": "all_done", "scan_id": scan_id})

    except Exception as e:
        logger.error("Scan thread failed: %s", e)
        _emit_event(scan_id, {"type": "error", "message": str(e)})
    finally:
        with _scan_lock:
            if scan_id in _active_scans:
                _active_scans[scan_id]["status"] = "done"


def _emit_event(scan_id: str, event: dict):
    with _scan_lock:
        if scan_id in _active_scans:
            _active_scans[scan_id]["events"].append(event)


# ──────────────────────────── HELPERS ────────────────────────────

def _sanitize_scan_id(scan_id: str) -> str | None:
    """Validate scan_id — prevent path traversal (allow only alphanumeric + - _)."""
    if not scan_id:
        return None
    safe = "".join(c for c in scan_id if c.isalnum() or c in "-_")
    if safe != scan_id or not safe:
        return None
    return safe


def _redact(value: str) -> str:
    """Redact API key for display — show only last 4 chars."""
    if not value:
        return ""
    if len(value) <= 8:
        return "****"
    return "****" + value[-4:]


# ─────────────────────────── ASSETS (DB-005) ──────────────────────────────────

@app.get("/api/assets")
async def list_assets():
    """Asset inventory: aggregate all discovered hosts from scan history."""
    assets: dict[str, dict] = {}
    scans_dir = Config.SCANS_DIR
    if not scans_dir.exists():
        return []
    for scan_dir in sorted(scans_dir.iterdir(), reverse=True):
        if not scan_dir.is_dir():
            continue
        data_file = scan_dir / "scan_data.json"
        if not data_file.exists():
            continue
        try:
            with open(data_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            target = data.get("target", "")
            if not target:
                continue
            risk_score = None
            analysis_file = scan_dir / "analysis.json"
            if analysis_file.exists():
                with open(analysis_file, "r", encoding="utf-8") as af:
                    analysis = json.load(af)
                risk_score = analysis.get("risk_score") or analysis.get("risicoscore")
            prev = assets.get(target, {})
            assets[target] = {
                "target": target,
                "target_type": data.get("target_type", "web"),
                "last_scan": data.get("started_at", ""),
                "last_scan_id": data.get("scan_id", scan_dir.name),
                "scan_count": prev.get("scan_count", 0) + 1,
                "risk_score": risk_score,
                "total_findings": data.get("total_findings", 0),
                "open_ports": _extract_open_ports(data),
                "os_guess": _extract_os(data),
            }
        except Exception:
            continue
    return sorted(assets.values(), key=lambda x: x.get("risk_score") or 0, reverse=True)


def _extract_open_ports(scan_data: dict) -> list[int]:
    ports: set[int] = set()
    for mod in scan_data.get("results", []):
        for finding in mod.get("findings", []):
            port = finding.get("port")
            if isinstance(port, int):
                ports.add(port)
    return sorted(ports)[:30]


def _extract_os(scan_data: dict) -> str:
    for mod in scan_data.get("results", []):
        for finding in mod.get("findings", []):
            if finding.get("type") == "os_detection":
                return finding.get("os", "")
    return ""


# ─────────────────────────── SCHEDULED SCANS (SE-004) ─────────────────────────

_schedules: list[dict] = []
_schedule_lock = threading.Lock()

try:
    from apscheduler.schedulers.background import BackgroundScheduler as _BgScheduler
    _scan_scheduler = _BgScheduler()
    _scan_scheduler.start()
except Exception:
    _scan_scheduler = None  # APScheduler not available


def _load_schedules_file() -> list[dict]:
    """Load schedules from data/schedules.json."""
    try:
        if Config.SCHEDULES_FILE.exists():
            with open(Config.SCHEDULES_FILE, "r", encoding="utf-8") as fh:
                return json.load(fh)
    except Exception as e:
        logger.warning("Could not load schedules file: %s", e)
    return []


def _save_schedules_file():
    """Persist _schedules to data/schedules.json."""
    try:
        Config.SCHEDULES_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(Config.SCHEDULES_FILE, "w", encoding="utf-8") as fh:
            json.dump(_schedules, fh, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning("Could not save schedules file: %s", e)


def _execute_scheduled_scan(schedule_id: str):
    """APScheduler callback: start a scan for the given schedule entry."""
    from datetime import timedelta
    with _schedule_lock:
        schedule = next((s for s in _schedules if s["id"] == schedule_id), None)
    if not schedule or not schedule.get("enabled"):
        return
    modules = Config.QUICK_MODULES[:] if schedule["scan_type"] == "quick" else Config.ALL_MODULES[:]
    scanner = Scanner(
        target=schedule["target"], modules=modules,
        scan_type=schedule["scan_type"], scan_mode=schedule["scan_mode"],
    )
    scan_id = scanner.scan_id
    with _scan_lock:
        _active_scans[scan_id] = {"status": "running", "events": []}
    thread = threading.Thread(target=_run_scan_thread, args=(scanner, scan_id), daemon=True)
    thread.start()
    now = datetime.now(timezone.utc)
    delta = {"daily": timedelta(days=1), "weekly": timedelta(weeks=1), "monthly": timedelta(days=30)}.get(
        schedule.get("interval", "weekly"), timedelta(weeks=1)
    )
    with _schedule_lock:
        for s in _schedules:
            if s["id"] == schedule_id:
                s["last_run"] = now.isoformat()
                s["next_run"] = (now + delta).isoformat()
    _save_schedules_file()
    _add_notification("Geplande scan gestart", f"Scan voor {schedule['target']} gestart", "info", scan_id)


def _register_schedule_job(schedule: dict):
    """Add or update an APScheduler interval job for a schedule entry."""
    if _scan_scheduler is None:
        return
    interval_days = {"daily": 1, "weekly": 7, "monthly": 30}.get(schedule.get("interval", "weekly"), 7)
    job_id = f"scan_schedule_{schedule['id']}"
    try:
        if _scan_scheduler.get_job(job_id):
            _scan_scheduler.remove_job(job_id)
        if schedule.get("enabled"):
            _scan_scheduler.add_job(
                _execute_scheduled_scan,
                trigger="interval",
                days=interval_days,
                id=job_id,
                args=[schedule["id"]],
                replace_existing=True,
            )
    except Exception as e:
        logger.warning("APScheduler job registration failed: %s", e)


class ScheduleRequest(BaseModel):
    target: str
    scan_type: str = "quick"
    scan_mode: str = "blackbox"
    target_type: str = "web"
    interval: str = "weekly"
    enabled: bool = True


@app.get("/api/schedules")
@app.get("/api/schedule")
async def list_schedules():
    with _schedule_lock:
        return list(_schedules)


@app.post("/api/schedules")
@app.post("/api/schedule")
async def create_schedule(body: ScheduleRequest):
    import uuid
    from datetime import timedelta
    now = datetime.now(timezone.utc)
    delta = {"daily": timedelta(days=1), "weekly": timedelta(weeks=1), "monthly": timedelta(days=30)}.get(body.interval, timedelta(weeks=1))
    schedule = {
        "id": uuid.uuid4().hex[:12],
        "target": body.target,
        "scan_type": body.scan_type,
        "scan_mode": body.scan_mode,
        "target_type": body.target_type,
        "interval": body.interval,
        "enabled": body.enabled,
        "created_at": now.isoformat(),
        "next_run": (now + delta).isoformat(),
        "last_run": None,
    }
    with _schedule_lock:
        _schedules.append(schedule)
    _register_schedule_job(schedule)
    _save_schedules_file()
    _add_audit_log("schedule_created", f"Schedule aangemaakt voor {body.target}")
    return schedule


@app.delete("/api/schedules/{schedule_id}")
async def delete_schedule(schedule_id: str):
    with _schedule_lock:
        for i, s in enumerate(_schedules):
            if s["id"] == schedule_id:
                _schedules.pop(i)
                if _scan_scheduler:
                    try:
                        _scan_scheduler.remove_job(f"scan_schedule_{schedule_id}")
                    except Exception:
                        pass
                _save_schedules_file()
                return {"status": "deleted"}
    raise HTTPException(status_code=404, detail="Schedule niet gevonden")


@app.patch("/api/schedules/{schedule_id}/toggle")
async def toggle_schedule(schedule_id: str, enabled: bool = True):
    with _schedule_lock:
        for s in _schedules:
            if s["id"] == schedule_id:
                s["enabled"] = enabled
                _register_schedule_job(s)
                _save_schedules_file()
                return s
    raise HTTPException(status_code=404)


# ─────────────────────────── FALSE POSITIVE MARKING (AI-001) ──────────────────

@app.post("/api/scans/{scan_id}/findings/{finding_idx}/mark")
async def mark_finding(scan_id: str, finding_idx: int, status: str = "confirmed"):
    """Mark a finding as confirmed, false_positive, or accepted_risk."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    if status not in ("confirmed", "false_positive", "accepted_risk"):
        raise HTTPException(status_code=400, detail="Ongeldige status")
    scan_dir = Config.SCANS_DIR / scan_id
    marks_file = scan_dir / "finding_marks.json"
    marks = {}
    if marks_file.exists():
        with open(marks_file, "r", encoding="utf-8") as f:
            marks = json.load(f)
    marks[str(finding_idx)] = {"status": status, "marked_at": datetime.now(timezone.utc).isoformat()}
    with open(marks_file, "w", encoding="utf-8") as f:
        json.dump(marks, f, ensure_ascii=False, indent=2)
    _add_audit_log("finding_marked", f"Bevinding {finding_idx} gemarkeerd als {status}")
    return {"status": "updated", "finding_idx": finding_idx, "mark": status}


@app.get("/api/scans/{scan_id}/findings/marks")
async def get_finding_marks(scan_id: str):
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    marks_file = Config.SCANS_DIR / scan_id / "finding_marks.json"
    if not marks_file.exists():
        return {}
    with open(marks_file, "r", encoding="utf-8") as f:
        return json.load(f)


# ─────────────────────────── AI CHAT (AI-005) ─────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    history: list[dict] | None = None


@app.post("/api/scans/{scan_id}/chat")
async def chat_with_scan(scan_id: str, body: ChatRequest):
    """Answer questions about a specific scan using AI."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    scan_dir = Config.SCANS_DIR / scan_id
    data_file = scan_dir / "scan_data.json"
    if not data_file.exists():
        raise HTTPException(status_code=404, detail="Scan niet gevonden")
    with open(data_file, "r", encoding="utf-8") as f:
        scan_data = json.load(f)
    analysis = {}
    analysis_file = scan_dir / "analysis.json"
    if analysis_file.exists():
        with open(analysis_file, "r", encoding="utf-8") as f:
            analysis = json.load(f)
    context = (
        f"Je bent een expert penetratietester. Scan context:\n"
        f"- Target: {scan_data.get('target', 'onbekend')}\n"
        f"- Type: {scan_data.get('scan_type', 'onbekend')}\n"
        f"- Bevindingen: {scan_data.get('total_findings', 0)}\n"
        f"- Risicoscore: {analysis.get('risicoscore', 'N/A')}\n"
        f"- Top 3 bevindingen: {str(analysis.get('bevindingen', [])[:3])}\n\n"
        f"Vraag: {body.message}\n\nAntwoord in het Nederlands, beknopt en professioneel."
    )
    try:
        from openai import OpenAI
        client = OpenAI(api_key=Config.DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
        messages = [{"role": "system", "content": "Je bent een expert penetratietester."}]
        messages += (body.history or [])[-6:]
        messages.append({"role": "user", "content": context})
        resp = client.chat.completions.create(model=Config.DEEPSEEK_MODEL, messages=messages, temperature=0.5, max_tokens=800)
        answer = resp.choices[0].message.content
    except Exception as e:
        logger.error("AI chat failed: %s", e)
        answer = "AI chat tijdelijk niet beschikbaar. Controleer je DeepSeek API key."
    chat_file = scan_dir / "chat_history.json"
    history = []
    if chat_file.exists():
        with open(chat_file, "r", encoding="utf-8") as f:
            history = json.load(f)
    ts = datetime.now(timezone.utc).isoformat()
    history.append({"role": "user", "content": body.message, "ts": ts})
    history.append({"role": "assistant", "content": answer, "ts": ts})
    with open(chat_file, "w", encoding="utf-8") as f:
        json.dump(history[-50:], f, ensure_ascii=False, indent=2)
    return {"answer": answer, "scan_id": scan_id}


@app.get("/api/scans/{scan_id}/chat/history")
async def get_chat_history(scan_id: str):
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    chat_file = Config.SCANS_DIR / scan_id / "chat_history.json"
    if not chat_file.exists():
        return []
    with open(chat_file, "r", encoding="utf-8") as f:
        return json.load(f)


# ─────────────────────────── DELTA SCAN / COMPARE (RP-011) ───────────────────

@app.get("/api/scans/compare")
async def compare_scans(scan_a: str, scan_b: str):
    """Compare two scans: new / resolved / unchanged findings."""
    scan_a = _sanitize_scan_id(scan_a)
    scan_b = _sanitize_scan_id(scan_b)
    if not scan_a or not scan_b:
        raise HTTPException(status_code=400, detail="Ongeldige scan IDs")

    def _load(sid: str) -> list[dict]:
        f = Config.SCANS_DIR / sid / "scan_data.json"
        if not f.exists():
            return []
        with open(f, "r", encoding="utf-8") as fh:
            d = json.load(fh)
        return [finding for mod in d.get("results", []) for finding in mod.get("findings", [])]

    def _key(f: dict) -> str:
        return f"{f.get('type', '')}-{f.get('target', f.get('url', ''))}-{f.get('port', '')}"

    fa = {_key(f): f for f in _load(scan_a)}
    fb = {_key(f): f for f in _load(scan_b)}
    return {
        "scan_a": scan_a, "scan_b": scan_b,
        "new": [f for k, f in fb.items() if k not in fa],
        "resolved": [f for k, f in fa.items() if k not in fb],
        "unchanged": [f for k, f in fb.items() if k in fa],
        "summary": {"new_count": len([k for k in fb if k not in fa]),
                    "resolved_count": len([k for k in fa if k not in fb]),
                    "unchanged_count": len([k for k in fb if k in fa])},
    }


# ─────────────────────────── WEBHOOKS (IN-003) ────────────────────────────────

_webhooks: list[dict] = []
_webhook_lock = threading.Lock()


class WebhookRequest(BaseModel):
    url: str
    events: list[str]
    enabled: bool = True


@app.get("/api/webhooks")
async def list_webhooks():
    with _webhook_lock:
        return [{"id": w["id"], "url": w["url"][:40] + "...", "events": w["events"], "enabled": w["enabled"]}
                for w in _webhooks]


@app.post("/api/webhooks")
async def create_webhook(body: WebhookRequest):
    import uuid, re
    if not re.match(r'^https://', body.url):
        raise HTTPException(status_code=400, detail="Alleen HTTPS URLs zijn toegestaan")
    wh = {"id": uuid.uuid4().hex[:12], "url": body.url, "events": body.events,
          "enabled": body.enabled, "created_at": datetime.now(timezone.utc).isoformat()}
    with _webhook_lock:
        _webhooks.append(wh)
    _add_audit_log("webhook_created", f"Webhook aangemaakt")
    return {"id": wh["id"], "url": body.url[:40] + "...", "events": body.events, "enabled": body.enabled}


@app.delete("/api/webhooks/{webhook_id}")
async def delete_webhook(webhook_id: str):
    with _webhook_lock:
        for i, w in enumerate(_webhooks):
            if w["id"] == webhook_id:
                _webhooks.pop(i)
                return {"status": "deleted"}
    raise HTTPException(status_code=404)


def _fire_webhooks(event_type: str, payload: dict):
    import urllib.request
    with _webhook_lock:
        active = [w for w in _webhooks if w["enabled"] and event_type in w["events"]]
    for wh in active:
        try:
            data = json.dumps({"event": event_type, "payload": payload, "ts": datetime.now(timezone.utc).isoformat()}).encode()
            req = urllib.request.Request(wh["url"], data=data, headers={"Content-Type": "application/json"}, method="POST")
            urllib.request.urlopen(req, timeout=5)
        except Exception as e:
            logger.warning("Webhook %s failed: %s", wh["url"][:30], e)


# ─────────────────────────── API KEYS (IN-002) ────────────────────────────────

import secrets as _secrets
_api_keys: list[dict] = []
_api_key_lock = threading.Lock()


class ApiKeyRequest(BaseModel):
    name: str
    scope: str = "read"


@app.get("/api/keys")
async def list_api_keys():
    with _api_key_lock:
        return [{"id": k["id"], "name": k["name"], "key": k["key"][:12] + "...",
                 "scope": k["scope"], "created_at": k["created_at"], "last_used": k["last_used"]}
                for k in _api_keys]


@app.post("/api/keys")
async def create_api_key(body: ApiKeyRequest):
    import uuid
    if body.scope not in ("read", "scan", "admin"):
        raise HTTPException(status_code=400, detail="Ongeldige scope")
    key = "cp_" + secrets.token_urlsafe(32)
    entry = {"id": uuid.uuid4().hex[:12], "name": body.name, "key": key,
              "scope": body.scope, "created_at": datetime.now(timezone.utc).isoformat(), "last_used": None}
    with _api_key_lock:
        _api_keys.append(entry)
    _add_audit_log("api_key_created", f"API key '{body.name}' aangemaakt")
    return entry


@app.delete("/api/keys/{key_id}")
async def delete_api_key(key_id: str):
    with _api_key_lock:
        for i, k in enumerate(_api_keys):
            if k["id"] == key_id:
                _api_keys.pop(i)
                return {"status": "deleted"}
    raise HTTPException(status_code=404)


# ─────────────────────────── AUDIT LOG (ST-009) ───────────────────────────────

_audit_log: list[dict] = []
_audit_lock = threading.Lock()


def _add_audit_log(action: str, detail: str = ""):
    entry = {"ts": datetime.now(timezone.utc).isoformat(), "action": action, "detail": detail}
    with _audit_lock:
        _audit_log.insert(0, entry)
        if len(_audit_log) > 1000:
            _audit_log.pop()


@app.get("/api/audit")
async def get_audit_log(limit: int = 100):
    with _audit_lock:
        return _audit_log[:limit]


# ─────────────────────────── NOTIFICATIONS (NT) ───────────────────────────────

_notifications: list[dict] = []
_notif_lock = threading.Lock()


def _add_notification(title: str, message: str, severity: str = "info", scan_id: str | None = None):
    import uuid
    entry = {"id": uuid.uuid4().hex[:12], "title": title, "message": message,
             "severity": severity, "scan_id": scan_id, "read": False,
             "created_at": datetime.now(timezone.utc).isoformat()}
    with _notif_lock:
        _notifications.insert(0, entry)
        if len(_notifications) > 200:
            _notifications.pop()


@app.get("/api/notifications")
async def list_notifications(limit: int = 30):
    with _notif_lock:
        return _notifications[:limit]


@app.get("/api/notifications/unread-count")
async def unread_count():
    with _notif_lock:
        return {"count": sum(1 for n in _notifications if not n["read"])}


@app.post("/api/notifications/{notif_id}/read")
async def mark_notification_read(notif_id: str):
    with _notif_lock:
        for n in _notifications:
            if n["id"] == notif_id:
                n["read"] = True
                return n
    raise HTTPException(status_code=404)


@app.post("/api/notifications/read-all")
async def mark_all_read():
    with _notif_lock:
        for n in _notifications:
            n["read"] = True
    return {"status": "ok"}


@app.delete("/api/notifications/all")
async def clear_notifications():
    with _notif_lock:
        _notifications.clear()
    return {"status": "cleared"}


# ─────────────────────────── AUTHORIZED SCAN (PS-010) ─────────────────────────

_authorizations: dict[str, dict] = {}
_auth_lock = threading.Lock()


class AuthorizationRequest(BaseModel):
    target: str
    authorized_by: str
    organization: str
    confirmed: bool


@app.post("/api/authorizations")
async def create_authorization(body: AuthorizationRequest):
    import uuid
    if not body.confirmed:
        raise HTTPException(status_code=400, detail="Bevestiging vereist")
    target = body.target.strip()
    if not target:
        raise HTTPException(status_code=400, detail="Target is verplicht")
    record = {"id": uuid.uuid4().hex[:12], "target": target, "authorized_by": body.authorized_by,
               "organization": body.organization, "confirmed": True,
               "created_at": datetime.now(timezone.utc).isoformat()}
    with _auth_lock:
        _authorizations[target] = record
    _add_audit_log("scan_authorized", f"Scan geautoriseerd voor {target}")
    return record


@app.get("/api/authorizations/{target:path}")
async def check_authorization(target: str):
    with _auth_lock:
        record = _authorizations.get(target)
    if not record:
        return {"authorized": False}
    return {"authorized": True, **record}


# ────────────── OWASP / CWE / MITRE / COMPLIANCE (RP-006, RP-018) ─────────────

OWASP_MAPPING: dict[str, dict] = {
    "sqli": {"owasp": "A03:2021", "cwe": "CWE-89", "label": "SQL Injection"},
    "xss": {"owasp": "A03:2021", "cwe": "CWE-79", "label": "Cross-Site Scripting"},
    "idor": {"owasp": "A01:2021", "cwe": "CWE-284", "label": "Broken Access Control"},
    "ssrf": {"owasp": "A10:2021", "cwe": "CWE-918", "label": "SSRF"},
    "xxe": {"owasp": "A05:2021", "cwe": "CWE-611", "label": "XXE Injection"},
    "open_redirect": {"owasp": "A01:2021", "cwe": "CWE-601", "label": "Open Redirect"},
    "sensitive_data": {"owasp": "A02:2021", "cwe": "CWE-312", "label": "Cryptographic Failure"},
    "auth_bypass": {"owasp": "A07:2021", "cwe": "CWE-287", "label": "Auth Failure"},
    "outdated_software": {"owasp": "A06:2021", "cwe": "CWE-1035", "label": "Vulnerable Component"},
    "misconfiguration": {"owasp": "A05:2021", "cwe": "CWE-16", "label": "Security Misconfiguration"},
    "lfi": {"owasp": "A03:2021", "cwe": "CWE-22", "label": "Path Traversal"},
    "rfi": {"owasp": "A03:2021", "cwe": "CWE-98", "label": "Remote File Inclusion"},
    "command_injection": {"owasp": "A03:2021", "cwe": "CWE-78", "label": "Command Injection"},
    "insecure_deserialization": {"owasp": "A08:2021", "cwe": "CWE-502", "label": "Insecure Deserialization"},
    "weak_password": {"owasp": "A07:2021", "cwe": "CWE-521", "label": "Weak Credentials"},
    "csrf": {"owasp": "A01:2021", "cwe": "CWE-352", "label": "CSRF"},
    "cors": {"owasp": "A05:2021", "cwe": "CWE-346", "label": "CORS Misconfiguration"},
    "ssl_tls": {"owasp": "A02:2021", "cwe": "CWE-326", "label": "TLS/SSL Issue"},
    "information_disclosure": {"owasp": "A05:2021", "cwe": "CWE-200", "label": "Information Disclosure"},
    "default_credentials": {"owasp": "A07:2021", "cwe": "CWE-1392", "label": "Default Credentials"},
    "directory_traversal": {"owasp": "A03:2021", "cwe": "CWE-22", "label": "Directory Traversal"},
    "prototype_pollution": {"owasp": "A08:2021", "cwe": "CWE-1321", "label": "Prototype Pollution"},
    "log4shell": {"owasp": "A06:2021", "cwe": "CWE-917", "label": "Log4Shell"},
    "deserializaton": {"owasp": "A08:2021", "cwe": "CWE-502", "label": "Deserialization"},
}

MITRE_MAPPING: dict[str, str] = {
    "sqli": "T1190", "xss": "T1059.007", "ssrf": "T1071.001",
    "command_injection": "T1059", "auth_bypass": "T1078", "weak_password": "T1110",
    "lfi": "T1083", "rfi": "T1105", "misconfiguration": "T1562",
    "outdated_software": "T1190", "information_disclosure": "T1530",
    "csrf": "T1185", "default_credentials": "T1078.001",
}

OWASP_TOP10 = [
    {"id": "A01:2021", "name": "Broken Access Control"},
    {"id": "A02:2021", "name": "Cryptographic Failures"},
    {"id": "A03:2021", "name": "Injection"},
    {"id": "A04:2021", "name": "Insecure Design"},
    {"id": "A05:2021", "name": "Security Misconfiguration"},
    {"id": "A06:2021", "name": "Vulnerable & Outdated Components"},
    {"id": "A07:2021", "name": "Identification & Authentication Failures"},
    {"id": "A08:2021", "name": "Software & Data Integrity Failures"},
    {"id": "A09:2021", "name": "Security Logging & Monitoring Failures"},
    {"id": "A10:2021", "name": "Server-Side Request Forgery"},
]


@app.get("/api/compliance/mapping")
async def get_compliance_mapping():
    return {"owasp": OWASP_MAPPING, "mitre": MITRE_MAPPING, "owasp_top10": OWASP_TOP10}


@app.get("/api/compliance/score/{scan_id}")
async def get_compliance_score(scan_id: str):
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    data_file = Config.SCANS_DIR / scan_id / "scan_data.json"
    if not data_file.exists():
        raise HTTPException(status_code=404)
    with open(data_file, "r", encoding="utf-8") as f:
        data = json.load(f)
    finding_types = {f.get("type", "").lower()
                     for mod in data.get("results", [])
                     for f in mod.get("findings", [])}
    owasp_found: dict[str, list] = {}
    for ft in finding_types:
        m = OWASP_MAPPING.get(ft)
        if m:
            owasp_found.setdefault(m["owasp"], []).append(
                {"type": ft, "cwe": m["cwe"], "label": m["label"],
                 "mitre": MITRE_MAPPING.get(ft, "")})
    return {
        "scan_id": scan_id,
        "owasp_findings": owasp_found,
        "owasp_categories_found": len(owasp_found),
        "owasp_categories_total": 10,
        "owasp_top10": OWASP_TOP10,
        "coverage_pct": round(len(owasp_found) / 10 * 100),
        "nis2_note": "NIS2 artikel 21 vereist aantoonbare technische beveiligingsmaatregelen.",
        "iso27001_note": "ISO 27001 Annex A.12: Operationele beveiliging & A.14: Systeemontwikkeling beveiliging.",
    }


# ─────────────────────────── WORD / CSV EXPORT (RP-002, RP-015) ───────────────

@app.get("/api/scans/{scan_id}/docx")
async def download_docx_report(scan_id: str):
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    scan_dir = Config.SCANS_DIR / scan_id
    data_file = scan_dir / "scan_data.json"
    if not data_file.exists():
        raise HTTPException(status_code=404, detail="Scan niet gevonden")
    with open(data_file, "r", encoding="utf-8") as f:
        scan_data = json.load(f)
    analysis: dict = {}
    analysis_file = scan_dir / "analysis.json"
    if analysis_file.exists():
        with open(analysis_file, "r", encoding="utf-8") as f:
            analysis = json.load(f)
    docx_path = scan_dir / "report.docx"
    try:
        _generate_docx(scan_data, analysis, str(docx_path))
        return FileResponse(str(docx_path),
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            filename=f"cyberpulse-{scan_id}.docx")
    except ImportError:
        # Fallback: plain text
        txt_path = scan_dir / "report.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"CyberPulse — Rapport\nTarget: {scan_data.get('target', '')}\n")
            f.write(f"Datum: {scan_data.get('started_at', '')}\nBevindingen: {scan_data.get('total_findings', 0)}\n\n")
            for mod in scan_data.get("results", []):
                for finding in mod.get("findings", []):
                    f.write(f"[{finding.get('severity', 'info').upper()}] {finding.get('title', finding.get('type', ''))}\n")
        return FileResponse(str(txt_path), media_type="text/plain", filename=f"cyberpulse-{scan_id}.txt")


def _generate_docx(scan_data: dict, analysis: dict, output_path: str):
    from docx import Document
    doc = Document()
    doc.add_heading("CyberPulse — Penetratietest Rapport", 0)
    doc.add_paragraph(f"Target: {scan_data.get('target', 'Onbekend')}")
    doc.add_paragraph(f"Datum: {scan_data.get('started_at', 'Onbekend')}")
    doc.add_paragraph(f"Risicoscore: {analysis.get('risicoscore', 'N/A')}/100")
    doc.add_heading("Management Samenvatting", 1)
    doc.add_paragraph(analysis.get("management_samenvatting", "Geen samenvatting beschikbaar."))
    doc.add_heading("Bevindingen", 1)
    for mod in scan_data.get("results", []):
        for finding in mod.get("findings", []):
            sev = finding.get("severity", "info").upper()
            doc.add_heading(f"[{sev}] {finding.get('title', finding.get('type', 'Bevinding'))}", 2)
            if finding.get("description"):
                doc.add_paragraph(finding["description"])
    doc.add_heading("Aanbevelingen", 1)
    for i, rec in enumerate(analysis.get("aanbevelingen_prioriteit", []), 1):
        doc.add_paragraph(f"{i}. {rec.get('titel', str(rec)) if isinstance(rec, dict) else rec}")
    doc.save(output_path)


@app.get("/api/scans/{scan_id}/csv")
async def download_csv_report(scan_id: str):
    """Export findings as CSV."""
    import csv, io
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    data_file = Config.SCANS_DIR / scan_id / "scan_data.json"
    if not data_file.exists():
        raise HTTPException(status_code=404)
    with open(data_file, "r", encoding="utf-8") as f:
        scan_data = json.load(f)
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["severity", "type", "title", "target", "port", "description"])
    for mod in scan_data.get("results", []):
        for finding in mod.get("findings", []):
            writer.writerow([
                finding.get("severity", "info"), finding.get("type", ""),
                finding.get("title", finding.get("titel", "")),
                finding.get("target", finding.get("url", "")),
                finding.get("port", ""), finding.get("description", finding.get("beschrijving", "")),
            ])
    csv_bytes = buf.getvalue().encode("utf-8")
    from fastapi.responses import Response
    return Response(content=csv_bytes, media_type="text/csv",
                    headers={"Content-Disposition": f"attachment; filename=cyberpulse-{scan_id}.csv"})


# ─────────────────────────── DEDUPLICATION (SE) ───────────────────────────────

@app.post("/api/scans/{scan_id}/deduplicate")
async def deduplicate_findings(scan_id: str):
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    data_file = Config.SCANS_DIR / scan_id / "scan_data.json"
    if not data_file.exists():
        raise HTTPException(status_code=404)
    with open(data_file, "r", encoding="utf-8") as f:
        scan_data = json.load(f)
    seen: set[str] = set()
    total_before = total_after = 0
    for mod in scan_data.get("results", []):
        deduped = []
        for finding in mod.get("findings", []):
            total_before += 1
            key = f"{finding.get('type', '')}-{finding.get('target', finding.get('url', ''))}-{finding.get('port', '')}"
            if key not in seen:
                seen.add(key)
                deduped.append(finding)
                total_after += 1
        mod["findings"] = deduped
    scan_data["total_findings"] = total_after
    with open(data_file, "w", encoding="utf-8") as f:
        json.dump(scan_data, f, ensure_ascii=False, indent=2)
    _add_audit_log("deduplication", f"{scan_id}: {total_before} → {total_after}")
    return {"before": total_before, "after": total_after, "removed": total_before - total_after}


# ─────────────────────────── SCAN PAUSE / RESUME (SE-009) ─────────────────────

@app.post("/api/scan/{scan_id}/pause")
async def pause_scan(scan_id: str):
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    with _scan_lock:
        state = _active_scans.get(scan_id)
        if not state:
            raise HTTPException(status_code=404)
        state["status"] = "paused"
        state["events"].append({"type": "scan_paused", "message": "Scan gepauzeerd"})
    return {"status": "paused"}


@app.post("/api/scan/{scan_id}/resume")
async def resume_scan(scan_id: str):
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    with _scan_lock:
        state = _active_scans.get(scan_id)
        if not state:
            raise HTTPException(status_code=404)
        state["status"] = "running"
        state["events"].append({"type": "scan_resumed", "message": "Scan hervat"})
    return {"status": "running"}


# ─────────────────────────── CI/CD ENDPOINT (DS-001) ──────────────────────────

class CiScanRequest(BaseModel):
    target: str
    scan_type: str = "quick"
    fail_on_critical: bool = True
    fail_on_high: bool = False
    fail_score_above: int | None = None


@app.post("/api/ci/scan")
async def ci_scan(body: CiScanRequest):
    """Trigger a scan from CI/CD pipeline. Returns pass/fail status."""
    target = body.target.strip()
    if not target:
        raise HTTPException(status_code=400, detail="Target is verplicht")
    from config import Config as C
    modules = C.QUICK_MODULES[:] if body.scan_type == "quick" else C.ALL_MODULES[:]
    scanner = Scanner(target=target, modules=modules, scan_type=body.scan_type)
    scan_id = scanner.scan_id
    with _scan_lock:
        _active_scans[scan_id] = {"status": "running", "events": []}
    results = list(scanner.run())
    scan_data_file = C.SCANS_DIR / scan_id / "scan_data.json"
    pass_scan = True
    reason = ""
    if scan_data_file.exists():
        with open(scan_data_file, "r", encoding="utf-8") as f:
            d = json.load(f)
        all_findings = [f for mod in d.get("results", []) for f in mod.get("findings", [])]
        critical_count = sum(1 for f in all_findings if f.get("severity") == "critical")
        high_count = sum(1 for f in all_findings if f.get("severity") == "high")
        if body.fail_on_critical and critical_count > 0:
            pass_scan = False
            reason = f"{critical_count} kritieke bevindingen gevonden"
        elif body.fail_on_high and high_count > 0:
            pass_scan = False
            reason = f"{high_count} hoge bevindingen gevonden"
    return {"scan_id": scan_id, "pass": pass_scan, "reason": reason, "status": "complete"}


# ─────────────────────────── THREAT INTELLIGENCE (AI-014) ─────────────────────

@app.get("/api/threat-intel/{cve_id}")
async def get_threat_intel(cve_id: str):
    """Get CVE threat intelligence from NVD API."""
    import re
    if not re.match(r'^CVE-\d{4}-\d+$', cve_id):
        raise HTTPException(status_code=400, detail="Ongeldig CVE formaat")
    try:
        import urllib.request
        url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?cveId={cve_id}"
        with urllib.request.urlopen(url, timeout=5) as resp:
            data = json.loads(resp.read())
        vuln = data.get("vulnerabilities", [{}])[0].get("cve", {})
        desc = next((d["value"] for d in vuln.get("descriptions", []) if d["lang"] == "en"), "")
        metrics = vuln.get("metrics", {})
        cvss = (metrics.get("cvssMetricV31") or metrics.get("cvssMetricV30") or [{}])[0]
        score = cvss.get("cvssData", {}).get("baseScore", "N/A")
        return {"cve_id": cve_id, "description": desc, "cvss_score": score, "source": "NVD"}
    except Exception as e:
        return {"cve_id": cve_id, "description": "Data tijdelijk niet beschikbaar", "cvss_score": "N/A", "error": str(e)}


# ─────────────────────────── GLOBAL AI CHAT (AI-005 global) ──────────────────

class GlobalChatRequest(BaseModel):
    message: str
    scan_id: str | None = None
    history: list[dict] | None = None


@app.post("/api/chat")
async def global_chat(body: GlobalChatRequest):
    """Chat with DeepSeek, optionally grounded in a specific scan's context."""
    context_lines: list[str] = []
    if body.scan_id:
        scan_id = _sanitize_scan_id(body.scan_id)
        if scan_id:
            data_file = Config.SCANS_DIR / scan_id / "scan_data.json"
            analysis_file = Config.SCANS_DIR / scan_id / "analysis.json"
            if data_file.exists():
                with open(data_file, "r", encoding="utf-8") as f:
                    sd = json.load(f)
                context_lines.append(f"Target: {sd.get('target', 'onbekend')}")
                context_lines.append(f"Scan type: {sd.get('scan_type', 'onbekend')}")
                context_lines.append(f"Totaal bevindingen: {sd.get('total_findings', 0)}")
            if analysis_file.exists():
                with open(analysis_file, "r", encoding="utf-8") as f:
                    an = json.load(f)
                context_lines.append(f"Risicoscore: {an.get('risicoscore', 'N/A')}")
                top3 = str(an.get("bevindingen", [])[:3])
                context_lines.append(f"Top bevindingen: {top3}")

    system_msg = "Je bent een expert penetratietester. Geef beknopte, professionele antwoorden in het Nederlands."
    user_content = body.message
    if context_lines:
        user_content = "Scan context:\n" + "\n".join(context_lines) + f"\n\nVraag: {body.message}"

    try:
        from openai import OpenAI as _OAI
        client = _OAI(api_key=Config.DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
        messages = [{"role": "system", "content": system_msg}]
        messages += (body.history or [])[-6:]
        messages.append({"role": "user", "content": user_content})
        resp = client.chat.completions.create(
            model=Config.DEEPSEEK_MODEL, messages=messages,
            temperature=0.5, max_tokens=800,
        )
        answer = resp.choices[0].message.content
    except Exception as e:
        logger.error("Global AI chat failed: %s", e)
        answer = "AI chat tijdelijk niet beschikbaar."

    return {"answer": answer}


# ─────────────────────────── COMPARE (path-param alias) ──────────────────────

@app.get("/api/compare/{scan_id_a}/{scan_id_b}")
async def compare_scans_path(scan_id_a: str, scan_id_b: str):
    """Delta compare — path-param version (mirrors /api/scans/compare query-param endpoint)."""
    return await compare_scans(scan_a=scan_id_a, scan_b=scan_id_b)


# ─────────────────────────── COMPLIANCE / NIS2 (RP-006 extended) ─────────────

# NIS2 article mapping onto OWASP categories
NIS2_MAPPING: dict[str, str] = {
    "A01:2021": "NIS2 Art. 21(2)(a) — Beleid voor risicoanalyse en informatiebeveiliging",
    "A02:2021": "NIS2 Art. 21(2)(h) — Beveiligde communicatie (versleuteling)",
    "A03:2021": "NIS2 Art. 21(2)(e) — Beveiliging bij verwerving, ontwikkeling en onderhoud",
    "A04:2021": "NIS2 Art. 21(2)(b) — Beveiliging van informatiesystemen (veilig ontwerp)",
    "A05:2021": "NIS2 Art. 21(2)(c) — Behandeling van incidenten (configuratiebeheer)",
    "A06:2021": "NIS2 Art. 21(2)(e) — Beheer van kwetsbaarheden in componenten",
    "A07:2021": "NIS2 Art. 21(2)(i) — Authenticatie en toegangsbeleid",
    "A08:2021": "NIS2 Art. 21(2)(e) — Integriteit van software en data",
    "A09:2021": "NIS2 Art. 21(2)(f) — Business continuity & crisismanagement (logging)",
    "A10:2021": "NIS2 Art. 21(2)(a) — Beveiliging van netwerken (SSRF / externe aanvragen)",
}

# ISO 27001 Annex A control mapping
ISO27001_MAPPING: dict[str, str] = {
    "A01:2021": "A.9 Toegangsbeveiliging",
    "A02:2021": "A.10 Cryptografie",
    "A03:2021": "A.14 Beveiliging van informatiesystemen",
    "A04:2021": "A.14.2 Beveiliging in ontwikkel- en ondersteuningsprocessen",
    "A05:2021": "A.12 Beveiliging betreffende bedrijfsvoering",
    "A06:2021": "A.12.6 Beheer van technische kwetsbaarheden",
    "A07:2021": "A.9.4 Toegangsbeveiliging op systeem- en applicatieniveau",
    "A08:2021": "A.14.2 Integriteitsborging software",
    "A09:2021": "A.12.4 Registratie en monitoring",
    "A10:2021": "A.13 Communicatiebeveiliging",
}


@app.get("/api/compliance/{scan_id}")
async def get_compliance_report(scan_id: str):
    """Full compliance report: OWASP, NIS2, ISO 27001, GDPR for a scan."""
    scan_id = _sanitize_scan_id(scan_id)
    if not scan_id:
        raise HTTPException(status_code=404)
    data_file = Config.SCANS_DIR / scan_id / "scan_data.json"
    if not data_file.exists():
        raise HTTPException(status_code=404, detail="Scan niet gevonden")
    with open(data_file, "r", encoding="utf-8") as f:
        scan_data = json.load(f)
    analysis: dict = {}
    analysis_file = Config.SCANS_DIR / scan_id / "analysis.json"
    if analysis_file.exists():
        with open(analysis_file, "r", encoding="utf-8") as f:
            analysis = json.load(f)

    # Collect all finding types from both scan_data and analysis
    finding_types: set[str] = set()
    for mod in scan_data.get("results", []):
        for finding in mod.get("findings", []):
            t = finding.get("type", "").lower().strip()
            if t:
                finding_types.add(t)
    for finding in analysis.get("bevindingen", []):
        t = finding.get("type", finding.get("categorie", "")).lower().strip()
        if t:
            finding_types.add(t)

    # Build per-OWASP-category detail with NIS2 + ISO annotations
    owasp_detail: dict[str, dict] = {}
    for ft in finding_types:
        m = OWASP_MAPPING.get(ft)
        if not m:
            continue
        cat_id = m["owasp"]
        if cat_id not in owasp_detail:
            owasp_cat = next((c for c in OWASP_TOP10 if c["id"] == cat_id), {})
            owasp_detail[cat_id] = {
                "id": cat_id,
                "name": owasp_cat.get("name", cat_id),
                "findings": [],
                "nis2": NIS2_MAPPING.get(cat_id, ""),
                "iso27001": ISO27001_MAPPING.get(cat_id, ""),
            }
        owasp_detail[cat_id]["findings"].append({
            "type": ft,
            "cwe": m["cwe"],
            "label": m["label"],
            "mitre": MITRE_MAPPING.get(ft, ""),
        })

    categories_found = len(owasp_detail)
    return {
        "scan_id": scan_id,
        "target": scan_data.get("target", ""),
        "scan_date": scan_data.get("started_at", ""),
        "risk_score": analysis.get("risicoscore"),
        "owasp_categories": owasp_detail,
        "owasp_categories_found": categories_found,
        "owasp_categories_total": 10,
        "coverage_pct": round(categories_found / 10 * 100),
        "owasp_top10": OWASP_TOP10,
        "nis2": {
            "applicable": categories_found > 0,
            "article": "NIS2 Richtlijn 2022/2555 — Artikel 21: Maatregelen voor cyberbeveiliging",
            "summary": (
                f"{categories_found} van 10 OWASP-categorieën aangetroffen. "
                "Organisaties die onder NIS2 vallen moeten passende technische maatregelen treffen."
            ),
            "per_category": {cat_id: NIS2_MAPPING.get(cat_id, "") for cat_id in owasp_detail},
        },
        "iso27001": {
            "applicable": categories_found > 0,
            "per_category": {cat_id: ISO27001_MAPPING.get(cat_id, "") for cat_id in owasp_detail},
        },
        "gdpr": {
            "note": (
                "GDPR Artikel 32 vereist passende technische en organisatorische maatregelen. "
                "Kwetsbaarheden gerelateerd aan A02 (cryptografie) en A07 (authenticatie) "
                "zijn direct relevant voor persoonsgegevens."
            ),
        },
    }


def _get_recent_scans(limit: int = 20) -> list[dict]:
    scans = []
    scans_dir = Config.SCANS_DIR
    if not scans_dir.exists():
        return scans

    for scan_dir in sorted(scans_dir.iterdir(), reverse=True):
        if not scan_dir.is_dir():
            continue
        data_file = scan_dir / "scan_data.json"
        if not data_file.exists():
            continue
        try:
            with open(data_file, "r", encoding="utf-8") as f:
                data = json.load(f)

            # Calculate risk score and severity counts from analysis if available
            risk_score = None
            severity_counts: dict = {}
            analysis_file = scan_dir / "analysis.json"
            if analysis_file.exists():
                try:
                    with open(analysis_file, "r", encoding="utf-8") as af:
                        analysis = json.load(af)
                    risk_score = analysis.get("risk_score") or analysis.get("risicoscore")
                    for f in analysis.get("bevindingen", []):
                        sev = f.get("ernst", "info")
                        severity_counts[sev] = severity_counts.get(sev, 0) + 1
                except (json.JSONDecodeError, OSError):
                    pass

            scans.append({
                "scan_id": data.get("scan_id", scan_dir.name),
                "target": data.get("target", ""),
                "scan_type": data.get("scan_type", ""),
                "scan_mode": data.get("scan_mode", "blackbox"),
                "target_type": data.get("target_type", "web"),
                "started_at": data.get("started_at", ""),
                "total_findings": data.get("total_findings", 0),
                "risk_score": risk_score,
                "severity_counts": severity_counts,
                "has_report": (scan_dir / "report.pdf").exists() or (scan_dir / "report.html").exists(),
                "has_analysis": analysis_file.exists(),
            })
        except (json.JSONDecodeError, OSError):
            continue

        if len(scans) >= limit:
            break

    return scans
